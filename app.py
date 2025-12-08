from flask import Flask, request, send_file, render_template
import os
import re
from collections import defaultdict
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

app = Flask(__name__)

UPLOAD_FOLDER = "uploads"
OUTPUT_FILE = "FINAL_DYNAMIC_TABLES.docx"
PDF_OUTPUT = "PDF_TO_DOCX_OUTPUT.docx"
GRADE_OUTPUT = "GRADE_PROCESSED.docx"

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)


# -------------------------------------------------------
# Add borders
# -------------------------------------------------------
def set_borders(table):
    tbl = table._element
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "10")
        elem.set(qn("w:color"), "000000")
        borders.append(elem)
    tbl.tblPr.append(borders)


# -------------------------------------------------------
# Header (logo + blue box)
# -------------------------------------------------------
def add_survey_header(doc, header_text):
    # Logo
    try:
        p_logo = doc.add_paragraph()
        p_logo.alignment = 1
        r = p_logo.add_run()
        r.add_picture("static/logo.jpg", width=Inches(2.8))
    except:
        print("Logo missing (static/logo.jpg)")

    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]

    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), "B7D3F2")
    cell._tc.get_or_add_tcPr().append(shade)
    set_borders(t)

    p = cell.paragraphs[0]
    p.alignment = 1
    for line in header_text.split("\n"):
        run = p.add_run(line)
        run.bold = True
        p.add_run("\n")

    doc.add_paragraph()


# -------------------------------------------------------
# UNIVERSAL TABLE HANDLER (for IG campuses)
# -------------------------------------------------------
def extract_table(tb, qnum, campus, data_dict):
    # Auto-assign default campus if missing
    if not campus:
        campus = "Percentage"

    header = [c.text.strip().lower() for c in tb.rows[0].cells]

    # Accept both "name" and "teacher"
    name_col = None
    percent_col = None
    ranking_col = None

    for idx, h in enumerate(header):
        if "name" in h or "teacher" in h:
            name_col = idx
        if "percentage" in h:
            percent_col = idx
        if "ranking" in h:
            ranking_col = idx

    is_ranking = ranking_col is not None
    actual_col = ranking_col if is_ranking else percent_col

    if name_col is None or actual_col is None:
        return False

    for row in tb.rows[1:]:
        name = row.cells[name_col].text.strip()

        if not name:
            continue

        # Normalize names
        name_normalized = name.strip().lower().title()

        raw = row.cells[actual_col].text.strip().replace("%", "")
        value = raw if is_ranking else (raw + "%" if raw else "")

        if value:
            data_dict[name_normalized][qnum][campus] = value

    return True



# -------------------------------------------------------
# Detect Campus Name (IG format)
# -------------------------------------------------------
def detect_campus(text):
    clean = " ".join(text.split())
    dash = r"[-–—]"

    patterns = [
        rf"(IG-[I1]+)\s*{dash}\s*(.+)$",
        rf"(IG-[I1]+)\s+(.+)$",
        rf"Grade(?:\s*\d+)?\s*{dash}\s*(.+)$",
        rf"Grade(?:\s*\d+)?\s+(.+)$",
    ]

    for pat in patterns:
        m = re.search(pat, clean, re.IGNORECASE)
        if m:
            if m.lastindex == 2:
                campus = m.group(2).strip()
            else:
                campus = m.group(1).strip()
                campus = re.sub(r"IG-[I1]+\s*", "", campus, flags=re.IGNORECASE).strip()
            
            return campus

    return None


# -------------------------------------------------------
# Process DOCX (IG campuses - existing logic)
# -------------------------------------------------------
def process_docx(path):
    doc = Document(path)
    paragraphs = doc.paragraphs
    tables = doc.tables

    teacher_data = defaultdict(lambda: defaultdict(dict))
    question_text = {}
    table_index = 0
    current_campus = None
    found_campuses = set()

    header_lines = []

    valid_keywords = [
        "learners",
        "academic year",
        "igcse boys",
        "igcse girls",
        "college campus",
        "igcse-i",
        "igcse ii",
        "igcse iii"
    ]

    # Extract header
    for p in paragraphs:
        t = p.text.strip()
        low = t.lower()

        if any(k in low for k in valid_keywords):
            header_lines.append(t)

        # Stop reading header once questions start
        if low.startswith("q#1") or low.startswith("q-1"):
            break

    survey_header = "\n".join(header_lines)

    # PARSE DOCUMENT
    for p in paragraphs:
        t = p.text.strip()

        # Detect campus (IG format)
        camp = detect_campus(t)
        if camp:
            current_campus = camp
            found_campuses.add(camp)

        # Detect questions
        m = re.match(r"(Q[#\-]\d+)", t, re.IGNORECASE)
        if m:
            qnum = m.group(1).upper().replace("-", "#")
            question_text[qnum] = t

            # ★ AUTO-ASSIGN CAMPUS USING Q#1 if none found ★
            if current_campus is None:
                current_campus = "Percentage"
                found_campuses.add("Percentage")

            # Extract next table
            while table_index < len(tables):
                tb = tables[table_index]
                table_index += 1
                if extract_table(tb, qnum, current_campus, teacher_data):
                    break

    # BUILD OUTPUT
    out = Document()
    campuses_sorted = sorted(found_campuses)

    for teacher, qs in teacher_data.items():
        add_survey_header(out, survey_header)
        out.add_heading(f"Teacher: {teacher}", level=2)

        # Reduce campuses to only those with data
        active_camps = [c for c in campuses_sorted if any(qs[q].get(c) for q in qs)]

        table = out.add_table(rows=1, cols=1 + len(active_camps))
        set_borders(table)

        hdr = table.rows[0].cells
        hdr[0].text = "Question"
        hdr[0].paragraphs[0].runs[0].bold = True

        for i, c in enumerate(active_camps):
            hdr[i + 1].text = c
            hdr[i + 1].paragraphs[0].runs[0].bold = True

        sorted_qs = sorted(qs.keys(), key=lambda x: int(re.findall(r"\d+", x)[0]))

        for q in sorted_qs:
            row = table.add_row().cells
            row[0].text = question_text[q]
            for i, c in enumerate(active_camps):
                row[i + 1].text = qs[q].get(c, "")

        # Monthly grading row
        # monthly_row = table.add_row().cells
        # monthly_row[0].text = "Monthly Grading"
        # monthly_row[0].paragraphs[0].runs[0].bold = True
        # for i in range(len(active_camps)):
        #     monthly_row[i + 1].text = ""

        # out.add_page_break()

    out.save(OUTPUT_FILE)
    return OUTPUT_FILE



# -------------------------------------------------------
# Process Grade-based DOCX (NEW FUNCTION)
# -------------------------------------------------------
# def process_grade_docx(path):
#     doc = Document(path)
#     paragraphs = doc.paragraphs
#     tables = doc.tables

#     teacher_data = defaultdict(lambda: defaultdict(dict))
#     question_text = {}
#     current_campus = None
#     found_campuses = []
#     table_index = 0

#     header_lines = []
#     valid_keywords = ["learners", "academic year", "pre-school", "preschool", "campus"]

#     # Extract header
#     for p in paragraphs:
#         t = p.text.strip()
#         low = t.lower()
        
#         if any(k in low for k in valid_keywords):
#             header_lines.append(t)
        
#         if low.startswith("q") or low.startswith("dated"):
#             break

#     survey_header = "\n".join(header_lines) if header_lines else "Learner's Survey\nAcademic Year 2025-2026"

#     # Parse document
#     for p in paragraphs:
#         t = p.text.strip()
#         low = t.lower()

#         # Detect campus (Grade 1 - Mars, Grade 1 - Venus, etc.)
#         grade_match = re.search(r"Grade\s+\d+\s*[-–—]\s*(.+)", t, re.IGNORECASE)
#         if grade_match:
#             current_campus = grade_match.group(1).strip()
#             found_campuses.append(current_campus)
#             continue

#         # Detect questions
#         q_match = re.match(r"(Q[-#]?\d+)[\:\.]?\s*(.+)", t, re.IGNORECASE)
#         if q_match and current_campus:
#             qnum = q_match.group(1).upper().replace("-", "#")
#             q_text = q_match.group(2).strip()
#             question_text[qnum] = f"{qnum}: {q_text}"

#             # Extract table data
#             if table_index < len(tables):
#                 tb = tables[table_index]
#                 table_index += 1

#                 # Parse table rows
#                 for row in tb.rows[1:]:  # Skip header
#                     try:
#                         cells = row.cells
#                         if len(cells) >= 3:
#                             teacher_name = cells[0].text.strip()
#                             percentage = cells[2].text.strip()
                            
#                             if teacher_name and percentage:
#                                 teacher_data[teacher_name][qnum][current_campus] = percentage
#                     except:
#                         continue

#     # Build output document
#     out = Document()
#     unique_campuses = []
#     for campus in found_campuses:
#         if campus not in unique_campuses:
#             unique_campuses.append(campus)

#     for teacher, qs in sorted(teacher_data.items()):
#         add_survey_header(out, survey_header)
#         out.add_heading(f"Teacher: {teacher}", level=2)

#         # Filter campuses where this teacher has data
#         active_camps = [c for c in unique_campuses if any(qs.get(q, {}).get(c) for q in qs)]

#         if not active_camps:
#             continue

#         # Create table
#         table = out.add_table(rows=1, cols=1 + len(active_camps))
#         set_borders(table)

#         # Headers
#         hdr = table.rows[0].cells
#         hdr[0].text = "Question"
#         hdr[0].paragraphs[0].runs[0].bold = True

#         for i, campus in enumerate(active_camps):
#             hdr[i + 1].text = campus
#             hdr[i + 1].paragraphs[0].runs[0].bold = True

#         # Questions
#         sorted_qs = sorted(qs.keys(), key=lambda x: int(re.findall(r"\d+", x)[0]) if re.findall(r"\d+", x) else 0)

#         for q in sorted_qs:
#             row = table.add_row().cells
#             row[0].text = question_text.get(q, q)
#             for i, campus in enumerate(active_camps):
#                 row[i + 1].text = qs[q].get(campus, "")

#         # ADD "Monthly Grading" ROW
#         monthly_row = table.add_row().cells
#         monthly_row[0].text = "Monthly Grading"
#         monthly_row[0].paragraphs[0].runs[0].bold = True
#         for i in range(len(active_camps)):
#             monthly_row[i + 1].text = ""

#         out.add_page_break()

#     out.save(GRADE_OUTPUT)
#     return GRADE_OUTPUT


# -------------------------------------------------------
# PDF → DOCX CONVERTER
# -------------------------------------------------------
def convert_pdf_to_docx(pdf_path, output_path="PDF_CONVERTED.docx", campus_name=""):
    import pdfplumber
    import re
    from collections import defaultdict
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def add_borders(table):
        tbl = table._element
        borders = OxmlElement('w:tblBorders')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            elem = OxmlElement(f'w:{edge}')
            elem.set(qn('w:val'), 'single')
            elem.set(qn('w:sz'), '10')
            elem.set(qn('w:color'), '000000')
            borders.append(elem)
        tbl.tblPr.append(borders)

    doc = Document()

    header_para = doc.add_paragraph()
    header_para.alignment = 1

    run1 = header_para.add_run("Learner's Survey\nAcademic Year 2025-2026\n")
    run1.bold = True

    if campus_name.strip():
        run2 = header_para.add_run(campus_name.strip())
    else:
        run2 = header_para.add_run("Campus name")

    run2.bold = True
    doc.add_paragraph()

    q_pattern = re.compile(r".*?(Q[#\s]*\d+)\s*[:\.\-]*\s*(.*)", re.IGNORECASE)
    teacher_pattern = re.compile(r"(.+?)\s+(\d+)$")
    ranking_pattern = re.compile(r"^\s*(\d+)\s+(.*)$")

    questions = {}
    current_q = None

    pdf = pdfplumber.open(pdf_path)

    for page in pdf.pages:
        text = page.extract_text() or ""

        for line in text.split("\n"):
            clean = line.strip()
            if not clean:
                continue

            mq = q_pattern.match(clean)
            if mq:
                q_id = mq.group(1).replace(" ", "").upper()
                q_text = mq.group(1) + " " + mq.group(2)

                current_q = q_id
                questions.setdefault(q_id, {"text": q_text, "entries": []})
                continue

            if not current_q:
                continue

            if current_q == "Q#8":
                mr = ranking_pattern.match(clean)
                if mr:
                    rank = mr.group(1)
                    teacher = mr.group(2).strip()
                    questions[current_q]["entries"].append((teacher, rank))
                continue

            mt = teacher_pattern.search(clean)
            if mt:
                teacher = mt.group(1).strip()
                count = int(mt.group(2))
                questions[current_q]["entries"].append((teacher, count))

    pdf.close()

    any_data = False

    for q_id, block in questions.items():
        entries = block["entries"]
        if not entries:
            continue

        any_data = True
        doc.add_heading(block["text"], level=2)

        if q_id == "Q#8":
            table = doc.add_table(rows=1, cols=2)
            add_borders(table)

            hdr = table.rows[0].cells
            hdr[0].text = "Teacher"
            hdr[1].text = "Ranking"

            for teacher, rank in entries:
                row = table.add_row().cells
                row[0].text = teacher
                row[1].text = str(rank)

            # ADD "Monthly Grading" ROW
            # monthly_row = table.add_row().cells
            # monthly_row[0].text = "Monthly Grading"
            # monthly_row[0].paragraphs[0].runs[0].bold = True
            # monthly_row[1].text = ""

            # doc.add_page_break()
            continue

        grouped = defaultdict(int)
        total = 0
        for teacher, count in entries:
            grouped[teacher] += count
            total += count

        table = doc.add_table(rows=1, cols=3)
        add_borders(table)

        hdr = table.rows[0].cells
        hdr[0].text = "Teacher"
        hdr[1].text = "Responses"
        hdr[2].text = "Percentage"

        sorted_teachers = sorted(
            grouped.items(),
            key=lambda x: (x[0].lower().startswith("none of the above"), x[0].lower())
)


        for teacher, count in sorted_teachers:
            row = table.add_row().cells
            row[0].text = teacher
            row[1].text = str(count)

            pct = round((count / total) * 100, 1) if total else 0
            row[2].text = f"{pct}%"

        # ADD "Monthly Grading" ROW
        # monthly_row = table.add_row().cells
        # monthly_row[0].text = "Monthly Grading"
        # monthly_row[0].paragraphs[0].runs[0].bold = True
        # monthly_row[1].text = ""
        # monthly_row[2].text = ""

        doc.add_page_break()

    if not any_data:
        raise Exception("Nothing detected in PDF.")

    doc.save(output_path)
    return output_path


# -------------------------------------------------------
# ROUTES
# -------------------------------------------------------
@app.route("/")
def index():
    return render_template("upload.html")


@app.route("/upload", methods=["POST"])
def upload():
    f = request.files["file"]
    file_path = os.path.join(UPLOAD_FOLDER, f.filename)
    f.save(file_path)
    output = process_docx(file_path)
    return send_file(output, as_attachment=True)


@app.route("/upload_grade", methods=["POST"])
def upload_grade():
    f = request.files["file"]
    file_path = os.path.join(UPLOAD_FOLDER, f.filename)
    f.save(file_path)
    output = process_grade_docx(file_path)
    return send_file(output, as_attachment=True)


@app.route("/convert_pdf", methods=["POST"])
def convert_pdf():
    f = request.files.get("pdf_file")
    campus_name = request.form.get("campus_name", "").strip()

    if not f:
        return "No file selected", 400

    pdf_path = os.path.join(UPLOAD_FOLDER, f.filename)
    f.save(pdf_path)

    try:
        output = convert_pdf_to_docx(pdf_path, campus_name=campus_name)
        return send_file(output, as_attachment=True)
    except Exception as e:
        return f"Error: {str(e)}", 500


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=True)